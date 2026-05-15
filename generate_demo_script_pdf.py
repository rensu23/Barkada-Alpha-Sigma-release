from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("deliverables")
DOCX_PATH = OUTPUT_DIR / "Barkada_Demo_Script.docx"


BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
SUBTLE = RGBColor(0x66, 0x66, 0x66)


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=BODY_COLOR):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_paragraph(paragraph, fill="F4F6F9"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    nf = normal.paragraph_format
    nf.space_before = Pt(0)
    nf.space_after = Pt(6)
    nf.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    tf = title.paragraph_format
    tf.space_before = Pt(0)
    tf.space_after = Pt(6)
    tf.line_spacing = 1.0

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, ACCENT),
        ("Heading 2", 13, 14, 7, ACCENT),
        ("Heading 3", 12, 10, 5, RGBColor(0x1F, 0x4D, 0x78)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.1


def add_body_paragraph(doc, text, bold_lead=None, italic=False):
    p = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(p, after=6, line=1.1)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        remainder = p.add_run(text[len(bold_lead):])
        set_run_font(remainder, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_quote_box(doc, text):
    p = doc.add_paragraph(style="Normal")
    shade_paragraph(p)
    set_paragraph_spacing(p, before=4, after=8, line=1.1)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(text)
    set_run_font(run, italic=True, color=RGBColor(0x33, 0x33, 0x33))


def build_document():
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph("Barkada Demo Script and Defense Notes", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(style="Normal")
    subtitle_run = subtitle.add_run(
        "Prepared as a simple, natural script for a first-year IT database project presentation."
    )
    set_run_font(subtitle_run, size=11, italic=True, color=SUBTLE)
    set_paragraph_spacing(subtitle, after=10, line=1.1)

    add_quote_box(
        doc,
        'Quick vibe for delivery: calm, respectful, and natural. Think "finals week but still trying to sound alive."',
    )

    doc.add_paragraph("Main Demo Script", style="Heading 1")
    add_body_paragraph(
        doc,
        "Good day, Sir/Ma'am. We are the group behind this project, and our system is designed to help groups manage contributions and track member payments in a more organized way.",
    )
    add_body_paragraph(
        doc,
        "Our project is useful for scenarios like class funds, event fees, or organization collections where a treasurer needs to monitor who has already paid, who is still pending, and what contribution is currently active.",
    )
    add_body_paragraph(
        doc,
        "For this demonstration, I will show the flow of the system from the landing page, to account access, to group creation, contribution setup, member joining, payment submission, and final payment confirmation.",
    )

    steps = [
        (
            "1. Landing page",
            "This is the landing page of our system. This serves as the entry point where users can understand the purpose of the platform and proceed to register or log in.",
        ),
        (
            "2. Register or log in",
            "Next, I will register a new account or log in using an existing one. The system checks the user information stored in the database and grants access once the credentials are valid.",
        ),
        (
            "3. Create a group as treasurer",
            "Now I am logged in as the treasurer. I will create a new group, for example BSIT 1A Class Fund. After creating the group, the system generates a unique join code that members can use to enter the correct group.",
        ),
        (
            "4. Create a contribution",
            "After that, I will add a sample contribution such as Class Fund, with a required amount and due date. This allows the treasurer to post a clear payment request that all members of the group can see.",
        ),
        (
            "5. Join as a member",
            "Next, I will switch to another account and log in as a member. Using the join code, the member can join the group without the treasurer manually encoding each person one by one.",
        ),
        (
            "6. Mark payment as pending",
            "Once the member opens the contribution, the payment can be marked as pending. This means the member has submitted the payment status, but it still needs verification from the treasurer before being finalized.",
        ),
        (
            "7. Confirm or reject payment",
            "Now I will return to the treasurer account. In the pending payment list, the treasurer can review the submission. If the payment is valid, the treasurer can confirm it as paid. If not, it can be rejected or left unconfirmed depending on the system flow.",
        ),
        (
            "8. Show updated payment status",
            "Finally, I will show the payment history or updated contribution status. We can see that the member's status has changed successfully, which proves that the system tracks payments properly from submission up to approval.",
        ),
    ]

    for heading, body in steps:
        doc.add_paragraph(heading, style="Heading 2")
        add_body_paragraph(doc, body)
        add_body_paragraph(doc, f'What to say if you want more detail: "{body}"', italic=True)

    doc.add_paragraph("Extra Lines If the Teacher Wants More Explanation", style="Heading 1")
    extra_lines = [
        "The system uses the database to store user accounts, group records, contribution details, and payment transactions in an organized way.",
        "Each action in the system updates related records, which shows how the application layer and the database layer work together.",
        "The pending-to-paid process is important because it reflects real-world verification. A member can submit a payment, but the treasurer still has to validate it before it becomes final.",
        "The system is designed to reduce manual tracking through chat messages, paper lists, or scattered spreadsheets.",
        "This project is not only about storing data. It is also about building a practical workflow that is useful in real group collection scenarios.",
    ]
    for line in extra_lines:
        add_body_paragraph(doc, line)

    doc.add_paragraph("If They Ask About Normalization", style="Heading 1")
    add_body_paragraph(
        doc,
        "A good short answer is: we normalized our database to reduce redundant data and avoid update inconsistencies. Instead of storing repeated information in one large table, we separated related data into different tables and connected them using keys.",
    )
    add_body_paragraph(
        doc,
        "You can also say: we mainly followed up to Third Normal Form. In First Normal Form, each field contains only one value and there are no repeating groups. In Second Normal Form, non-key attributes depend on the whole primary key. In Third Normal Form, non-key attributes depend only on the primary key and not on other non-key attributes.",
    )
    add_body_paragraph(
        doc,
        "Example explanation: instead of repeating user details in every payment record, we store the user information once in the users table, then use user_id as a foreign key in related tables.",
    )

    doc.add_paragraph("Tables You Can Mention", style="Heading 2")
    table_lines = [
        "users - stores account information",
        "groups - stores the created groups",
        "group_members - connects users to groups",
        "contributions - stores payment requests such as class fund or event fee",
        "payments - stores payment status such as pending or paid",
    ]
    for line in table_lines:
        add_body_paragraph(doc, line)

    doc.add_paragraph("If They Ask About Keys", style="Heading 2")
    add_body_paragraph(
        doc,
        "A primary key uniquely identifies each record in a table, like user_id or group_id. A foreign key connects one table to another, such as group_id in contributions or user_id in payments. These relationships help maintain referential integrity in the database.",
    )

    doc.add_paragraph("If They Ask What Problem the System Solves", style="Heading 2")
    add_body_paragraph(
        doc,
        "Our system solves the problem of manual contribution tracking. Normally, treasurers rely on chats, notes, or spreadsheets, which can become confusing and inconsistent. Our platform centralizes contribution posting, joining, payment monitoring, and payment confirmation in one system.",
    )

    doc.add_paragraph("Short Closing Statement", style="Heading 1")
    add_body_paragraph(
        doc,
        "In summary, our project applies database concepts in a practical system where users, groups, contributions, and payments are all connected properly. Through this workflow, the system becomes more organized, trackable, and useful for real group payment scenarios. Thank you.",
    )

    doc.add_paragraph("Very Short Emergency Version", style="Heading 1")
    add_body_paragraph(
        doc,
        "Good day, Sir/Ma'am. Our project is a contribution and payment tracking system. The treasurer can create a group, post a contribution, and monitor member payments. The member can join through a code, view the contribution, and submit payment as pending. After verification, the treasurer confirms it as paid. This proves that the system can properly track collections in an organized database-driven workflow.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
